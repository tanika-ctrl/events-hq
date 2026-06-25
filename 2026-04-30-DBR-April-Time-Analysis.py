#!/usr/bin/env python3
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Styles
doc.styles['Normal'].font.name = 'Arial'
doc.styles['Normal'].font.size = Pt(11)
doc.styles['Heading 1'].font.name = 'Arial'
doc.styles['Heading 1'].font.size = Pt(16)
doc.styles['Heading 1'].font.bold = True
doc.styles['Heading 2'].font.name = 'Arial'
doc.styles['Heading 2'].font.size = Pt(13)
doc.styles['Heading 2'].font.bold = True

def shade_cell(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def set_borders(cell):
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_element = OxmlElement('w:' + edge)
        edge_element.set(qn('w:val'), 'single')
        edge_element.set(qn('w:sz'), '4')
        edge_element.set(qn('w:space'), '0')
        edge_element.set(qn('w:color'), 'CCCCCC')
        tcBorders.append(edge_element)
    tcPr.append(tcBorders)

# Header
header_para = doc.add_paragraph()
header_para.paragraph_format.space_after = Pt(12)
pPr = header_para._element.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'thick')
bottom.set(qn('w:sz'), '24')
bottom.set(qn('w:space'), '1')
bottom.set(qn('w:color'), '1a1a1a')
pBdr.append(bottom)
pPr.append(pBdr)

run1 = header_para.add_run('DBR — April 2026 | Time & Activity Analysis')
run1.font.bold = True
run1.font.size = Pt(11)
run1.font.name = 'Arial'
header_para.add_run('\t')
run2 = header_para.add_run('Prepared by Subsurface | 30 April 2026 | Source: Timesheet v009')
run2.font.size = Pt(11)
run2.font.name = 'Arial'

# Summary
doc.add_heading('Summary', level=1)
table1 = doc.add_table(rows=5, cols=3)

for i, label in enumerate(['Metric', 'Value', 'Note']):
    cell = table1.rows[0].cells[i]
    shade_cell(cell, '222222')
    set_borders(cell)
    cell.text = label
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(11)
            run.font.name = 'Arial'

summary_data = [
    ['Total hours logged', '244', 'April 2026 · DBR'],
    ['Team members', '4', 'BH · RK · SS · CJ'],
    ['Top work type', '3D work', '53h — OG48 model'],
    ['Entries logged', '84', 'Across 20 working days'],
]

for row_idx, row_data in enumerate(summary_data):
    row = table1.rows[row_idx + 1]
    for col_idx, text in enumerate(row_data):
        cell = row.cells[col_idx]
        if row_idx % 2 == 0:
            shade_cell(cell, 'f7f7f7')
        set_borders(cell)
        cell.text = text
        for para in cell.paragraphs:
            for run in para.runs:
                if col_idx == 1:
                    run.font.bold = True
                    run.font.size = Pt(20)
                elif col_idx == 2:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(102, 102, 102)
                else:
                    run.font.size = Pt(11)
                run.font.name = 'Arial'

doc.add_paragraph()
doc.add_heading('Hours by Work Type', level=2)

table2 = doc.add_table(rows=10, cols=3)
for i, label in enumerate(['Work Type', 'Hours', '% of Month']):
    cell = table2.rows[0].cells[i]
    shade_cell(cell, '222222')
    set_borders(cell)
    cell.text = label
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(11)
            run.font.name = 'Arial'

hours_data = [
    ['3D work', '53.0h', '21.7%'],
    ['Travel', '45.0h', '18.4%'],
    ['Office / yard', '43.0h', '17.6%'],
    ['Meeting', '36.5h', '14.9%'],
    ['Admin', '32.0h', '13.1%'],
    ['Research / R&D', '13.5h', '5.5%'],
    ['Event (IGSHPA)', '12.0h', '4.9%'],
    ['Video editing', '5.0h', '2.0%'],
    ['Project management', '4.0h', '1.6%'],
]

for row_idx, row_data in enumerate(hours_data):
    row = table2.rows[row_idx + 1]
    for col_idx, text in enumerate(row_data):
        cell = row.cells[col_idx]
        if row_idx % 2 == 0:
            shade_cell(cell, 'f7f7f7')
        set_borders(cell)
        cell.text = text
        for para in cell.paragraphs:
            if col_idx > 0:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in para.runs:
                run.font.size = Pt(11)
                run.font.name = 'Arial'

total_row = table2.rows[9]
for col_idx, text in enumerate(['TOTAL', '244.0h', '100%']):
    cell = total_row.cells[col_idx]
    shade_cell(cell, '222222')
    set_borders(cell)
    cell.text = text
    for para in cell.paragraphs:
        if col_idx > 0:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(11)
            run.font.name = 'Arial'

doc.add_paragraph()
doc.add_heading('Team Breakdown by Work Type', level=2)

table3 = doc.add_table(rows=10, cols=7)
for i, label in enumerate(['Work Type', 'BH', 'SS', 'RK', 'CJ', 'Total', '%']):
    cell = table3.rows[0].cells[i]
    shade_cell(cell, '222222')
    set_borders(cell)
    cell.text = label
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)
            run.font.name = 'Arial'

team_data = [
    ['3D work', '53.0h', '—', '—', '—', '53.0h', '21.7%'],
    ['Travel', '45.0h', '—', '—', '—', '45.0h', '18.4%'],
    ['Office / yard', '—', '9.0h', '12.0h', '22.0h', '43.0h', '17.6%'],
    ['Meeting', '10.75h', '15.75h', '10.0h', '—', '36.5h', '14.9%'],
    ['Admin', '4.5h', '6.5h', '13.0h', '8.0h', '32.0h', '13.1%'],
    ['Research / R&D', '9.5h', '4.0h', '—', '—', '13.5h', '5.5%'],
    ['Event (IGSHPA)', '12.0h', '—', '—', '—', '12.0h', '4.9%'],
    ['Video editing', '5.0h', '—', '—', '—', '5.0h', '2.0%'],
    ['Project mgmt', '—', '4.0h', '—', '—', '4.0h', '1.6%'],
]

for row_idx, row_data in enumerate(team_data):
    row = table3.rows[row_idx + 1]
    for col_idx, text in enumerate(row_data):
        cell = row.cells[col_idx]
        if row_idx % 2 == 0:
            shade_cell(cell, 'f7f7f7')
        set_borders(cell)
        cell.text = text
        for para in cell.paragraphs:
            if col_idx > 0:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.name = 'Arial'

total_row = table3.rows[9]
for col_idx, text in enumerate(['Total', '139.75h', '39.25h', '35.0h', '30.0h', '244.0h', '100%']):
    cell = total_row.cells[col_idx]
    shade_cell(cell, '222222')
    set_borders(cell)
    cell.text = text
    for para in cell.paragraphs:
        if col_idx > 0:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)
            run.font.name = 'Arial'

doc.add_paragraph()
doc.add_heading('April Highlights', level=2)

highlights = [
    ('OG48 3D Modelling', '53 hours across the month — BH\'s primary output, with near-daily sessions building out the OG48 rig model. The concentrated effort suggests a client-facing deadline on the horizon.'),
    ('IGSHPA Conference — Missouri', 'BH attended 13–16 April. 40 hours of travel and 12 hours at the conference and networking. Key industry presence and likely prospect conversations.'),
    ('The Yard — Office Build-Out', 'Approximately 51 hours combined across CJ, RK, and SS on physical office setup, sorting, and operations. Significant overhead but a one-time investment as the space becomes operational.'),
    ('New Relationships Activated', 'Danny Faubert appeared on 22 April (DBR App and manuals). Canadian partners had a 3-hour session each with RK and SS on 24 April. Jeff is a recurring contact throughout.'),
    ('Loom Video Recordings', '5 hours of Loom content produced by BH on 23 and 28 April — likely sales or client documentation for Salesio and new contacts. Signals the product is moving toward a demonstration stage.'),
    ('R&D — Ongoing Baseline', '13.5 hours across BH and SS. BH maintained near-daily 0.5–2h R&D sessions throughout the month; SS ran a concentrated 4h block on 8 April focused on DBR-specific research.'),
]

for title, text in highlights:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(title)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(12)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Arial'

doc.add_paragraph()
doc.add_heading('Analysis', level=2)

analysis = [
    'BH carried the majority of logged hours in April at 139.75h (57% of the month). Strip out 52 hours of travel and conference attendance for Missouri and he logged approximately 88 hours of productive output — a demanding month, anchored almost entirely by the OG48 3D model.',
    'The Yard is quietly consuming meaningful team capacity. Between CJ, RK, and SS, roughly 51 hours went into physically setting up and running the DBR office space across Office and Admin entries. This is operational overhead that will normalise once the space is fully functional, but it\'s worth tracking as it\'s currently displacing more strategic output — particularly from RK, whose April was largely admin, meetings, and yard time compared to his heavier strategic involvement in February and March.',
    'SS split her time between commercial meetings (15.75h), R&D, logistics, and office work. Her involvement in the Canadian partner conversations on 24 April and the Bruce/Jeff discussions in early April suggests she remains the key relationship lead on the commercial side.',
    'CJ\'s activity was entirely operational — 30 hours split between yard setup and admin. No meetings, design, or strategic involvement logged in April.',
    'Looking at the shape of the month: the IGSHPA Conference, new contacts (Danny Faubert, Canadians), and the progression of Loom recordings all suggest DBR is moving from development into active commercialisation. The OG48 model being the dominant deliverable of the month supports this — it\'s the asset that sits at the centre of the sales story.',
]

for text in analysis:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(12)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Arial'

# Add footer
section = doc.sections[0]
footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.text = 'DBR — April 2026 Time Analysis | Source: Subsurface Timesheet v009 | Prepared 30 April 2026'
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in footer_para.runs:
    run.font.size = Pt(10)
    run.font.name = 'Arial'

doc.save('/Users/tanikafacey/Desktop/Cowork/2026-04-30-DBR-April-Time-Analysis.docx')
print('Document created successfully')
